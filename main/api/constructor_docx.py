from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, Cm
import pymorphy2
import os


def _get_month(month):
    """Получить название месяца через номер."""
    names = {
        1: 'января',
        2: 'февраля',
        3: 'марта',
        4: 'апреля',
        5: 'мая',
        6: 'июня',
        7: 'июля',
        8: 'августа',
        9: 'сентября',
        10: 'октября',
        11: 'ноября',
        12: 'декабря'
    }
    return names[month]


def _get_date_live(date):
    """Получить название следующего месяца исходя из данного номера месяца."""
    return f'1 {_get_month(date.month + 1)}'


def _get_yaer_live_end(date):
    """Получить года учёбы курса."""
    if date.month >= 1 and date.month <= 7:
        return date.year
    else:
        return '{date.year}-{date.year + 1}'


def _text_children(children):
    """ПОлучить подходящий текст к слову РЕБЕНОК исходя из числа."""
    if children >= 10 and children <= 20:
        return 'детей'
    else:
        children %= 10
        if children == 1:
            return 'ребенок'
        elif children >= 2 and children <= 4:
            return 'ребенка'
        else:
            return 'детей'


def direction(
    name_f,
    name_l,
    dorm,
    address,
    phone,
    patronymic=None
):
    """Создание направления в виде документа под расширением .docx."""
    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(1.7)
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.7)
        section.right_margin = Cm(1.7)

    doc.add_paragraph()
    doc.add_paragraph()

    name = 'НАПРАВЛЕНИЕ'
    t = 'о предоставлении общежития обучающемуся в'
    university = 'Карагандинском экономическом университете Казпотребсоюза, по адресу г.Караганда, ул.Академическая, 9'
    doc.add_paragraph(f'{name}\n{t}\n{university}')
    doc.paragraphs[2].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    doc.add_paragraph(
        f'Гражданину {pymorphy2.MorphAnalyzer().parse(name_l)[0].inflect({"datv"}).word.title()} ' +
        f'{pymorphy2.MorphAnalyzer().parse(name_f)[0].inflect({"datv"}).word.title()} ' +
        f'{pymorphy2.MorphAnalyzer().parse(patronymic)[0].inflect({"datv"}).word.title() if patronymic != None else ""}'
    )

    dorm_number = f'Направление на заселение в общежитие № {dorm}'
    address_student = f'Адрес общежития {address}'
    doc.add_paragraph(f'{dorm_number}\n{address_student}')

    doc.add_paragraph()
    doc.add_paragraph()

    paragraph = doc.add_paragraph(
        'Ректор ______________\n    (подпись)\n«___» _____________ 20___ г.\n    М.П.')

    paragraph.style.font.size = Pt(16)
    paragraph.paragraph_format.line_spacing = 1
    doc.styles['Normal'].font.name = 'Times New Roman'

    if not os.path.isdir('docs'):
        os.mkdir('docs')

    doc.save(f'docs/direction_{phone}.docx')

    return f'docs/direction_{phone}.docx'


def request(
    dorm,
    date,
    number_request,
    gender,
    address,
    name_f,
    name_l,
    phone,
    group='ERROR-ERROR',
    children=1,
    patronymic=None
):
    """Создание заявления в виде документа под расширением .docx."""
    group = group if group != None else 'ERROR-ERROR'
    children = children if children != None else 1
    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(1.7)
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.7)
        section.right_margin = Cm(1.7)

    doc.add_paragraph(f'№-{number_request}')
    doc.paragraphs[0].runs[0].font.size = Pt(10)
    doc.paragraphs[0].runs[0].font.bold = True

    one = 'Ректору Карагандинского экономического универстета'
    two = 'д.э.н., профессору Е.Б.Аймагамбетову'
    three = f"Студента {pymorphy2.MorphAnalyzer().parse(name_l)[0].inflect({'gent'}).word.title()} {name_f[0]}.{patronymic[0] + '.' if patronymic != None else ''}"
    four = f'Группы {group}'
    five = f'Конт.тел. {phone}'
    doc.add_paragraph().add_run(
        f'{one}\n{two}\n{three}\n{four}\n{five}')
    doc.paragraphs[1].runs[0].font.bold = True
    doc.paragraphs[1].paragraph_format.left_indent = Inches(4.0)

    doc.add_paragraph()

    paragraph = doc.add_paragraph().add_run('ЗАЯВЛЕНИЕ')
    doc.paragraphs[3].runs[0].font.bold = True
    doc.paragraphs[3].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Прошу Вас предоставить мне место в общежитии ' +
                      f'№ {dorm} c {_get_date_live(date)} на {_get_yaer_live_end(date)} учебный год. ' +
                      f'В семье {children} {_text_children(children)}.')
    doc.paragraphs[4].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.paragraphs[4].paragraph_format.first_line_indent = Inches(0.43)

    doc.add_paragraph()

    doc.add_paragraph(f'Адрес проживания {address}')
    doc.paragraphs[6].paragraph_format.first_line_indent = Inches(0.43)

    doc.add_paragraph()

    doc.add_paragraph(
        'Своевременную оплату за проживание в общежитии ' +
        'и соблюдения порядка гарантирую. С «Правилами проживания в общежитии КЭУК» ' +
        f'{"ознакомлен и согласен" if gender == 1 else "ознакомлена и согласена"}.'
    )
    doc.paragraphs[8].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.paragraphs[8].paragraph_format.first_line_indent = Inches(0.43)

    doc.add_paragraph()
    paragraph = doc.add_paragraph()

    table = doc.add_table(rows=2, cols=2)
    row = table.rows[0]
    row.cells[0].text = 'Дата'
    row.cells[1].text = f'{date.day} {_get_month(date.month)} {date.year}г.'
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row = table.rows[1]
    row.cells[0].text = 'ФИО и подпись'
    row.cells[1].text = f'{name_l} {name_f} {patronymic if patronymic != None else ""}'
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    paragraph.style.font.size = Pt(16)
    doc.styles['Normal'].font.name = 'Times New Roman'
    paragraph.paragraph_format.line_spacing = 1

    if not os.path.isdir('docs'):
        os.mkdir('docs')

    doc.save(f'docs/request_{phone}.docx')

    return f'docs/request_{phone}.docx'
